"""
Gemma Studio — local Gemma 3 desktop chat (CustomTkinter UI).

No Ollama. Downloads the Gemma 3 GGUF (+ mmproj vision encoder for 4b/12b/27b)
from Hugging Face, auto-launches the bundled llama.cpp server (./llama_bin, or
./llama_bin_cuda when GPU is on) on 127.0.0.1, and streams over localhost.
Runs 100% offline on CPU after the first download.

Features:
  - Modern CustomTkinter UI with a conversations sidebar (SQLite-backed).
  - New / switch / rename / delete conversations, persisted across runs.
  - Model picker (1b/4b/12b/27b), GPU toggle (warns on driver fallback),
    image attach (vision), per-response metrics, live CPU/RAM.

Run:  .venv\Scripts\python.exe chat_app.py   (or double-click run.bat)
"""

import os
import re
import json
import base64
import queue
import socket
import sqlite3
import subprocess
import threading
import time
import datetime
import urllib.request
import urllib.error
import tkinter as tk
from tkinter import filedialog, messagebox

import psutil
import customtkinter as ctk
from PIL import Image
from huggingface_hub import hf_hub_download

# --- Config -----------------------------------------------------------------
DEFAULT_SIZE = "4b"

_MODELS = {
    "1b":  {"repo": "unsloth/gemma-3-1b-it-GGUF",  "file": "gemma-3-1b-it-Q4_K_M.gguf",  "mmproj": None,             "gpu_layers": 99},
    "4b":  {"repo": "unsloth/gemma-3-4b-it-GGUF",  "file": "gemma-3-4b-it-Q4_K_M.gguf",  "mmproj": "mmproj-F16.gguf", "gpu_layers": 99},
    "12b": {"repo": "unsloth/gemma-3-12b-it-GGUF", "file": "gemma-3-12b-it-Q4_K_M.gguf", "mmproj": "mmproj-F16.gguf", "gpu_layers": 16},
    "27b": {"repo": "unsloth/gemma-3-27b-it-GGUF", "file": "gemma-3-27b-it-Q4_K_M.gguf", "mmproj": "mmproj-F16.gguf", "gpu_layers": 8},
}

SYSTEM_PROMPT = "You are a helpful, friendly assistant running fully offline on the user's machine."
N_CTX = 8192
N_THREADS = os.cpu_count() or 4
HERE = os.path.dirname(os.path.abspath(__file__))
SERVER_EXE_CPU = os.path.join(HERE, "llama_bin", "llama-server.exe")
SERVER_EXE_CUDA = os.path.join(HERE, "llama_bin_cuda", "llama-server.exe")
DB_PATH = os.path.join(HERE, "chats.db")
ICON_ICO = os.path.join(HERE, "assets", "icon.ico")
ICON_PNG = os.path.join(HERE, "assets", "icon.png")
CREATE_NO_WINDOW = 0x08000000
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"}
ATTACH_TYPES = [
    ("Supported files", "*.png *.jpg *.jpeg *.bmp *.gif *.webp *.pdf *.docx *.txt *.md "
                        "*.csv *.json *.py *.js *.ts *.html *.css *.log *.xml *.yaml *.yml"),
    ("All files", "*.*"),
]
MAX_DOC_CHARS = 12000   # cap extracted document text so it fits the context window

# --- Palette (Catppuccin Mocha) --------------------------------------------
BG       = "#1e1e2e"
SIDEBAR  = "#181825"
CARD     = "#11111b"
SURFACE  = "#313244"
SURFACE2 = "#45475a"
TEXT     = "#cdd6f4"
SUBTEXT  = "#a6adc8"
MUTED    = "#6c7086"
BLUE     = "#89b4fa"
BLUE_HV  = "#74a8f0"
MAUVE    = "#cba6f7"
GREEN    = "#a6e3a1"
TEAL     = "#94e2d5"
PINK     = "#f5c2e7"
YELLOW   = "#f9e2af"
RED      = "#f38ba8"
RED_HV   = "#eb6f8e"


def _now():
    return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _free_port():
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.bind(("127.0.0.1", 0))
    port = s.getsockname()[1]
    s.close()
    return port


# ============================================================================
# SQLite persistence
# ============================================================================
class ChatStore:
    def __init__(self, path):
        self.conn = sqlite3.connect(path)
        self.conn.execute("PRAGMA journal_mode=WAL")
        self.conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS chats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL, model TEXT,
                created_at TEXT, updated_at TEXT
            );
            CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id INTEGER NOT NULL, role TEXT NOT NULL,
                content TEXT, image_path TEXT, created_at TEXT,
                FOREIGN KEY(chat_id) REFERENCES chats(id) ON DELETE CASCADE
            );
            """
        )
        self.conn.commit()

    def create_chat(self, model, title="New conversation"):
        cur = self.conn.execute(
            "INSERT INTO chats(title, model, created_at, updated_at) VALUES(?,?,?,?)",
            (title, model, _now(), _now()))
        self.conn.commit()
        return cur.lastrowid

    def list_chats(self):
        return self.conn.execute(
            "SELECT id, title, model, updated_at FROM chats ORDER BY updated_at DESC").fetchall()

    def add_message(self, chat_id, role, content, image_path=None):
        self.conn.execute(
            "INSERT INTO messages(chat_id, role, content, image_path, created_at) VALUES(?,?,?,?,?)",
            (chat_id, role, content, image_path, _now()))
        self.conn.execute("UPDATE chats SET updated_at=? WHERE id=?", (_now(), chat_id))
        self.conn.commit()

    def get_messages(self, chat_id):
        return self.conn.execute(
            "SELECT role, content, image_path FROM messages WHERE chat_id=? ORDER BY id",
            (chat_id,)).fetchall()

    def message_count(self, chat_id):
        return self.conn.execute("SELECT COUNT(*) FROM messages WHERE chat_id=?", (chat_id,)).fetchone()[0]

    def title_of(self, chat_id):
        r = self.conn.execute("SELECT title FROM chats WHERE id=?", (chat_id,)).fetchone()
        return r[0] if r else "Conversation"

    def set_title(self, chat_id, title):
        self.conn.execute("UPDATE chats SET title=? WHERE id=?", (title, chat_id)); self.conn.commit()

    def set_model(self, chat_id, model):
        self.conn.execute("UPDATE chats SET model=? WHERE id=?", (model, chat_id)); self.conn.commit()

    def delete_chat(self, chat_id):
        self.conn.execute("DELETE FROM messages WHERE chat_id=?", (chat_id,))
        self.conn.execute("DELETE FROM chats WHERE id=?", (chat_id,)); self.conn.commit()


# ============================================================================
# Main application
# ============================================================================
class ChatApp:
    def __init__(self, root):
        self.root = root
        self.store = ChatStore(DB_PATH)

        self.current_size = DEFAULT_SIZE
        self.cfg = _MODELS[self.current_size]
        self.supports_vision = self.cfg["mmproj"] is not None

        self.proc = None
        self.port = None
        self.base_url = None
        self.current_chat_id = None
        self.chat_buttons = {}
        self.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        self.pending_attach = None          # {"path","name","kind":"image"|"doc","text"?}
        self.events = queue.Queue()
        self._ps = {}
        self.use_gpu_var = tk.BooleanVar(value=False)

        ctk.set_appearance_mode("dark")
        self.root.title("Gemma Studio")
        self.root.geometry("1080x720")
        self.root.minsize(900, 580)
        self.root.configure(fg_color=BG)
        try:
            self.root.iconbitmap(ICON_ICO)
        except Exception:
            pass

        self._fonts()
        self._build_ui()
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        self.root.after(50, self._drain_events)
        self.root.after(1000, self._monitor)

        self._refresh_chat_list()
        self._new_chat()
        threading.Thread(target=self._start_engine, daemon=True).start()

    def _fonts(self):
        self.f_brand = ctk.CTkFont("Segoe UI Semibold", 17)
        self.f_h = ctk.CTkFont("Segoe UI Semibold", 14)
        self.f_label = ctk.CTkFont("Segoe UI", 13)
        self.f_small = ctk.CTkFont("Segoe UI", 11)
        self.f_tiny = ctk.CTkFont("Segoe UI", 10, "bold")
        self.f_btn = ctk.CTkFont("Segoe UI Semibold", 13)
        self.f_mono = ctk.CTkFont("Consolas", 12)

    # --- UI ----------------------------------------------------------------
    def _build_ui(self):
        self.root.grid_columnconfigure(1, weight=1)
        self.root.grid_rowconfigure(0, weight=1)

        # ===== Sidebar =====
        side = ctk.CTkFrame(self.root, width=262, corner_radius=0, fg_color=SIDEBAR)
        side.grid(row=0, column=0, rowspan=2, sticky="nsew")
        side.grid_propagate(False)
        side.grid_rowconfigure(3, weight=1)
        side.grid_columnconfigure(0, weight=1)

        try:
            logo = ctk.CTkImage(Image.open(ICON_PNG), size=(26, 26))
        except Exception:
            logo = None
        brand = ctk.CTkLabel(side, text="  Gemma Studio", image=logo, compound="left",
                             font=self.f_brand, text_color=TEXT, anchor="w")
        brand.grid(row=0, column=0, sticky="ew", padx=18, pady=(18, 12))

        self.new_btn = ctk.CTkButton(side, text="✚   New conversation", command=self._new_chat,
                                     font=self.f_btn, fg_color=BLUE, hover_color=BLUE_HV,
                                     text_color=BG, height=40, corner_radius=10, anchor="w")
        self.new_btn.grid(row=1, column=0, sticky="ew", padx=16, pady=(0, 14))

        ctk.CTkLabel(side, text="CONVERSATIONS", font=self.f_tiny, text_color=MUTED,
                     anchor="w").grid(row=2, column=0, sticky="ew", padx=20, pady=(0, 2))

        self.chat_scroll = ctk.CTkScrollableFrame(side, fg_color="transparent")
        self.chat_scroll.grid(row=3, column=0, sticky="nsew", padx=8)
        self.chat_scroll.grid_columnconfigure(0, weight=1)

        bottom = ctk.CTkFrame(side, fg_color="transparent")
        bottom.grid(row=4, column=0, sticky="ew", padx=14, pady=12)
        bottom.grid_columnconfigure((0, 1), weight=1)
        self.rename_btn = ctk.CTkButton(bottom, text="✎  Rename", command=self._rename_chat,
                                        font=self.f_small, fg_color=SURFACE, hover_color=SURFACE2,
                                        text_color=SUBTEXT, height=32, corner_radius=8)
        self.rename_btn.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.delete_btn = ctk.CTkButton(bottom, text="🗑  Delete", command=self._delete_chat,
                                        font=self.f_small, fg_color=SURFACE, hover_color=RED_HV,
                                        text_color=RED, height=32, corner_radius=8)
        self.delete_btn.grid(row=0, column=1, sticky="ew", padx=(4, 0))

        # ===== Main =====
        main = ctk.CTkFrame(self.root, corner_radius=0, fg_color=BG)
        main.grid(row=0, column=1, sticky="nsew")
        main.grid_columnconfigure(0, weight=1)
        main.grid_rowconfigure(2, weight=1)

        header = ctk.CTkFrame(main, fg_color="transparent")
        header.grid(row=0, column=0, sticky="ew", padx=18, pady=(16, 6))
        header.grid_columnconfigure(5, weight=1)
        ctk.CTkLabel(header, text="Model", font=self.f_label, text_color=SUBTEXT).grid(row=0, column=0, padx=(0, 6))
        self.model_menu = ctk.CTkOptionMenu(header, values=list(_MODELS.keys()), command=self._on_model_change,
                                            width=88, font=self.f_label, fg_color=SURFACE, button_color=SURFACE2,
                                            button_hover_color=BLUE, text_color=TEXT, corner_radius=8)
        self.model_menu.set(self.current_size)
        self.model_menu.grid(row=0, column=1)
        self.gpu_switch = ctk.CTkSwitch(header, text="GPU", variable=self.use_gpu_var, command=self._on_gpu_toggle,
                                        font=self.f_label, progress_color=GREEN, text_color=TEXT)
        self.gpu_switch.grid(row=0, column=2, padx=(16, 0))
        self.save_btn = ctk.CTkButton(header, text="💾  Save reply", command=self._on_save, width=116, height=30,
                                      font=self.f_small, fg_color=SURFACE, hover_color=SURFACE2,
                                      text_color=TEXT, corner_radius=8)
        self.save_btn.grid(row=0, column=3, padx=(16, 0))
        self.res_label = ctk.CTkLabel(header, text="CPU —   RAM —", font=self.f_mono, text_color=TEAL)
        self.res_label.grid(row=0, column=5, sticky="e")

        self.chat_title = ctk.CTkLabel(main, text="", font=self.f_h, text_color=TEXT, anchor="w")
        self.chat_title.grid(row=1, column=0, sticky="ew", padx=20, pady=(2, 6))

        # transcript: tk.Text inside a rounded card (reliable tag styling + streaming)
        card = ctk.CTkFrame(main, corner_radius=14, fg_color=CARD)
        card.grid(row=2, column=0, sticky="nsew", padx=16, pady=(0, 8))
        card.grid_rowconfigure(0, weight=1)
        card.grid_columnconfigure(0, weight=1)
        self.txt = tk.Text(card, wrap="word", bg=CARD, fg=TEXT, relief="flat", bd=0,
                           padx=18, pady=16, font=("Segoe UI", 12), highlightthickness=0,
                           insertbackground=TEXT, state="disabled", spacing3=3)
        self.txt.grid(row=0, column=0, sticky="nsew", padx=(6, 0), pady=6)
        sb = ctk.CTkScrollbar(card, command=self.txt.yview, button_color=SURFACE2, button_hover_color=MUTED)
        sb.grid(row=0, column=1, sticky="ns", padx=(0, 4), pady=8)
        self.txt.configure(yscrollcommand=sb.set)
        self.txt.tag_config("user", foreground=BLUE, font=("Segoe UI Semibold", 12), spacing1=8)
        self.txt.tag_config("ai", foreground=GREEN, font=("Segoe UI Semibold", 12), spacing1=8)
        self.txt.tag_config("sys", foreground=MUTED, font=("Segoe UI", 10, "italic"))
        self.txt.tag_config("img", foreground=PINK, font=("Segoe UI", 11, "italic"), spacing1=8)
        self.txt.tag_config("body", foreground=TEXT, lmargin1=4, lmargin2=4)
        self.txt.tag_config("metrics", foreground=MUTED, font=("Consolas", 10))
        self.txt.tag_config("warn", foreground=YELLOW, font=("Segoe UI", 10, "italic"))
        # markdown tags
        self.txt.tag_config("mdbold", foreground=TEXT, font=("Segoe UI Semibold", 12))
        self.txt.tag_config("mditalic", font=("Segoe UI", 12, "italic"))
        self.txt.tag_config("mdcode", font=("Consolas", 11), background=SURFACE, foreground=PINK)
        self.txt.tag_config("mdcodeblock", font=("Consolas", 11), background=SIDEBAR, foreground=GREEN,
                            lmargin1=18, lmargin2=18, spacing1=2, spacing3=2)
        self.txt.tag_config("mdh1", foreground=TEXT, font=("Segoe UI Semibold", 16), spacing1=10, spacing3=4)
        self.txt.tag_config("mdh2", foreground=TEXT, font=("Segoe UI Semibold", 14), spacing1=8, spacing3=3)
        self.txt.tag_config("mdh3", foreground=BLUE, font=("Segoe UI Semibold", 12), spacing1=6, spacing3=2)
        self.txt.tag_config("mdbullet", foreground=BLUE, font=("Segoe UI", 12), lmargin1=12, lmargin2=28)

        self.attach_label = ctk.CTkLabel(main, text="", font=self.f_small, text_color=PINK, anchor="w")
        self.attach_label.grid(row=3, column=0, sticky="ew", padx=20)

        bar = ctk.CTkFrame(main, fg_color="transparent")
        bar.grid(row=4, column=0, sticky="ew", padx=16, pady=(2, 12))
        bar.grid_columnconfigure(1, weight=1)
        self.attach_btn = ctk.CTkButton(bar, text="📎", command=self._on_attach, width=46, height=42,
                                        font=ctk.CTkFont(size=18), fg_color=SURFACE, hover_color=SURFACE2,
                                        text_color=TEXT, corner_radius=10)
        self.attach_btn.grid(row=0, column=0, padx=(0, 8))
        self.entry = ctk.CTkEntry(bar, placeholder_text="Type a message…", font=self.f_label,
                                  height=42, corner_radius=12, fg_color=SURFACE, border_width=0, text_color=TEXT)
        self.entry.grid(row=0, column=1, sticky="ew")
        self.entry.bind("<Return>", lambda e: self._on_send())
        self.send_btn = ctk.CTkButton(bar, text="Send  ➤", command=self._on_send, width=96, height=42,
                                      font=self.f_btn, fg_color=BLUE, hover_color=BLUE_HV, text_color=BG,
                                      corner_radius=12)
        self.send_btn.grid(row=0, column=2, padx=(8, 0))

        self.status = ctk.CTkLabel(self.root, text="Starting…", font=self.f_small, text_color=YELLOW,
                                   anchor="w", fg_color=CARD, height=26)
        self.status.grid(row=1, column=1, sticky="ew")

        self._apply_vision_ui()
        self._set_input_enabled(False)

    def _apply_vision_ui(self):
        self.attach_btn.configure(state="normal")  # files always allowed; image vision checked at attach time

    def _set_input_enabled(self, enabled):
        state = "normal" if enabled else "disabled"
        for w in (self.entry, self.send_btn, self.model_menu, self.gpu_switch,
                  self.new_btn, self.rename_btn, self.delete_btn):
            w.configure(state=state)
        for b in self.chat_buttons.values():
            b.configure(state=state)
        self.attach_btn.configure(state=state)
        if enabled:
            self.entry.focus_set()

    def _append(self, tag, label, text):
        self.txt.configure(state="normal")
        if label:
            self.txt.insert("end", label, tag)
        self.txt.insert("end", text, "body")
        self.txt.see("end")
        self.txt.configure(state="disabled")

    def _clear_display(self):
        self.txt.configure(state="normal")
        self.txt.delete("1.0", "end")
        self.txt.configure(state="disabled")

    # --- Markdown rendering (caller manages the widget state) --------------
    _INLINE = re.compile(r"(\*\*.+?\*\*|__.+?__|`[^`]+`|\*[^*\n]+\*|_[^_\n]+_)")

    def _insert_inline(self, text):
        pos = 0
        for mm in self._INLINE.finditer(text):
            if mm.start() > pos:
                self.txt.insert("end", text[pos:mm.start()], "body")
            tok = mm.group()
            if tok.startswith("**") or tok.startswith("__"):
                self.txt.insert("end", tok[2:-2], ("body", "mdbold"))
            elif tok.startswith("`"):
                self.txt.insert("end", tok[1:-1], ("body", "mdcode"))
            else:
                self.txt.insert("end", tok[1:-1], ("body", "mditalic"))
            pos = mm.end()
        if pos < len(text):
            self.txt.insert("end", text[pos:], "body")

    def _insert_markdown(self, md):
        in_code = False
        for line in (md or "").split("\n"):
            s = line.lstrip()
            if s.startswith("```"):
                in_code = not in_code
                continue
            if in_code:
                self.txt.insert("end", line + "\n", ("mdcodeblock",))
                continue
            if s.startswith("### "):
                self.txt.insert("end", s[4:] + "\n", ("mdh3",)); continue
            if s.startswith("## "):
                self.txt.insert("end", s[3:] + "\n", ("mdh2",)); continue
            if s.startswith("# "):
                self.txt.insert("end", s[2:] + "\n", ("mdh1",)); continue
            mb = re.match(r"^(\s*)[\*\-]\s+(.*)$", line)
            if mb:
                self.txt.insert("end", mb.group(1) + "•  ", ("mdbullet",))
                self._insert_inline(mb.group(2))
                self.txt.insert("end", "\n"); continue
            mn = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
            if mn:
                self.txt.insert("end", mn.group(1) + mn.group(2) + ".  ", ("mdbullet",))
                self._insert_inline(mn.group(3))
                self.txt.insert("end", "\n"); continue
            self._insert_inline(line)
            self.txt.insert("end", "\n")

    def _finalize_reply(self, reply, metrics):
        """Replace the streamed plain text of the last reply with formatted markdown."""
        self.txt.configure(state="normal")
        self.txt.delete(self._ai_body_index, "end")
        self.txt.insert("end", "\n")
        self._insert_markdown(reply)
        self.txt.insert("end", "\n")
        self.txt.insert("end", metrics, "metrics")
        self.txt.see("end")
        self.txt.configure(state="disabled")

    # --- Conversations -----------------------------------------------------
    def _refresh_chat_list(self):
        for b in self.chat_buttons.values():
            b.destroy()
        self.chat_buttons = {}
        for i, (cid, title, model, updated) in enumerate(self.store.list_chats()):
            label = title if len(title) <= 26 else title[:25] + "…"
            selected = (cid == self.current_chat_id)
            b = ctk.CTkButton(
                self.chat_scroll, text="  " + label, anchor="w", font=self.f_label,
                fg_color=SURFACE2 if selected else "transparent", hover_color=SURFACE,
                text_color=TEXT if selected else SUBTEXT, corner_radius=8, height=34,
                command=lambda c=cid: self._load_chat(c))
            b.grid(row=i, column=0, sticky="ew", pady=2, padx=2)
            self.chat_buttons[cid] = b

    def _highlight(self, cid):
        for c, b in self.chat_buttons.items():
            sel = (c == cid)
            b.configure(fg_color=SURFACE2 if sel else "transparent", text_color=TEXT if sel else SUBTEXT)

    def _new_chat(self):
        cid = self.store.create_chat(self.current_size)
        self.current_chat_id = cid
        self.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        self._clear_display()
        self.chat_title.configure(text="New conversation")
        self._clear_attach()
        self._refresh_chat_list()
        hint = "Start chatting"
        hint += " — you can attach an image too." if self.supports_vision else "."
        self._append("sys", "", hint + "\n\n")

    def _load_chat(self, cid):
        if cid == self.current_chat_id:
            return
        self.current_chat_id = cid
        self._highlight(cid)
        self.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        self._clear_display()
        self.chat_title.configure(text=self.store.title_of(cid))
        for role, content, image_path in self.store.get_messages(cid):
            if role == "user":
                if image_path:
                    is_img = self._is_image(image_path)
                    name = os.path.basename(image_path)
                    self._append("img", "You", f"   [{'🖼' if is_img else '📎'} {name}] ")
                    self._append("body", "", (content or "") + "\n\n")
                    if os.path.exists(image_path):
                        if is_img:
                            self.messages.append({"role": "user", "content": self._image_content(content, image_path)})
                        else:
                            self.messages.append({"role": "user", "content": self._doc_content(name, self._extract_text(image_path), content)})
                    else:
                        self.messages.append({"role": "user", "content": content or ""})
                else:
                    self._append("user", "You", "\n")
                    self._append("body", "", (content or "") + "\n\n")
                    self.messages.append({"role": "user", "content": content})
            elif role == "assistant":
                self.txt.configure(state="normal")
                self.txt.insert("end", "Gemma\n", "ai")
                self._insert_markdown(content or "")
                self.txt.insert("end", "\n\n")
                self.txt.configure(state="disabled")
                self.txt.see("end")
                self.messages.append({"role": "assistant", "content": content})

    def _rename_chat(self):
        if self.current_chat_id is None:
            return
        dlg = ctk.CTkInputDialog(text="New conversation name:", title="Rename")
        new = dlg.get_input()
        if new and new.strip():
            self.store.set_title(self.current_chat_id, new.strip())
            self.chat_title.configure(text=new.strip())
            self._refresh_chat_list()

    def _delete_chat(self):
        if self.current_chat_id is None:
            return
        if not messagebox.askyesno("Delete conversation", "Delete this conversation permanently?"):
            return
        self.store.delete_chat(self.current_chat_id)
        self.current_chat_id = None
        remaining = self.store.list_chats()
        if remaining:
            self._refresh_chat_list()
            self._load_chat(remaining[0][0])
        else:
            self._new_chat()

    # --- Attachments (images via vision, documents via text extraction) ----
    @staticmethod
    def _is_image(path):
        return os.path.splitext(path)[1].lower() in IMAGE_EXTS

    def _on_attach(self):
        path = filedialog.askopenfilename(title="Attach a file", filetypes=ATTACH_TYPES)
        if not path:
            return
        name = os.path.basename(path)
        if self._is_image(path):
            if not self.supports_vision:
                self._append("warn", "", f"\n[attach] Gemma 3 {self.current_size} can't see images — "
                                         f"switch to 4b/12b/27b, or attach a document instead.\n\n")
                return
            self.pending_attach = {"path": path, "name": name, "kind": "image"}
            self.attach_label.configure(text=f"🖼 {name}   (click to remove)")
        else:
            text = self._extract_text(path)
            self.pending_attach = {"path": path, "name": name, "kind": "doc", "text": text}
            self.attach_label.configure(text=f"📎 {name}  ·  {len(text):,} chars extracted   (click to remove)")
        self.attach_label.bind("<Button-1>", lambda e: self._clear_attach())

    def _clear_attach(self):
        self.pending_attach = None
        self.attach_label.configure(text="")
        self.attach_label.unbind("<Button-1>")

    def _extract_text(self, path):
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".pdf":
                from pypdf import PdfReader
                text = "\n".join((pg.extract_text() or "") for pg in PdfReader(path).pages)
            elif ext == ".docx":
                import docx
                text = "\n".join(p.text for p in docx.Document(path).paragraphs)
            else:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        except Exception as e:
            return f"(could not read {os.path.basename(path)}: {e})"
        if len(text) > MAX_DOC_CHARS:
            text = text[:MAX_DOC_CHARS] + "\n\n[... document truncated to fit the context window ...]"
        return text

    @staticmethod
    def _image_content(question, path):
        return [
            {"type": "text", "text": question or "Describe this image."},
            {"type": "image_url", "image_url": {"url": ChatApp._image_data_uri(path)}},
        ]

    @staticmethod
    def _doc_content(name, text, question):
        return (f"[Attached document: {name}]\n\n{text}\n\n---\n\n"
                f"{question or 'Please review the attached document.'}")

    @staticmethod
    def _image_data_uri(path):
        ext = os.path.splitext(path)[1].lstrip(".").lower() or "png"
        if ext == "jpg":
            ext = "jpeg"
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        return f"data:image/{ext};base64,{b64}"

    def _on_save(self):
        last = next((m["content"] for m in reversed(self.messages)
                     if m["role"] == "assistant" and isinstance(m["content"], str) and m["content"].strip()), None)
        if not last:
            messagebox.showinfo("Save reply", "There's no response to save yet.")
            return
        path = filedialog.asksaveasfilename(
            title="Save reply", defaultextension=".md", initialfile="gemma_reply",
            filetypes=[("Markdown", "*.md"), ("Text", "*.txt"), ("Word document", "*.docx")])
        if not path:
            return
        try:
            if path.lower().endswith(".docx"):
                import docx
                d = docx.Document()
                for para in last.split("\n"):
                    d.add_paragraph(para)
                d.save(path)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(last)
            self.status.configure(text=f"  Saved → {os.path.basename(path)}", text_color=GREEN)
        except Exception as e:
            messagebox.showerror("Save failed", str(e))

    # --- Model / GPU switching ---------------------------------------------
    def _on_model_change(self, size):
        if size == self.current_size:
            return
        self.current_size = size
        self.cfg = _MODELS[size]
        self.supports_vision = self.cfg["mmproj"] is not None
        if self.current_chat_id is not None:
            self.store.set_model(self.current_chat_id, size)
        self._clear_attach()
        self._apply_vision_ui()
        self._set_input_enabled(False)
        self._append("sys", "", f"\n— Switching to Gemma 3 {size} —\n\n")
        threading.Thread(target=self._start_engine, daemon=True).start()

    def _on_gpu_toggle(self):
        if self.use_gpu_var.get() and not os.path.exists(SERVER_EXE_CUDA):
            self.use_gpu_var.set(False)
            self._append("warn", "", "\n[GPU] CUDA build not found in llama_bin_cuda — staying on CPU.\n\n")
            return
        mode = "GPU" if self.use_gpu_var.get() else "CPU"
        self._set_input_enabled(False)
        self._append("sys", "", f"\n— Restarting engine in {mode} mode —\n\n")
        threading.Thread(target=self._start_engine, daemon=True).start()

    # --- Engine ------------------------------------------------------------
    def _start_engine(self):
        try:
            self._stop_server()
            self.port = _free_port()
            self.base_url = f"http://127.0.0.1:{self.port}"

            self.events.put(("status", f"Downloading {self.cfg['file']} (first run only)…"))
            model_path = hf_hub_download(repo_id=self.cfg["repo"], filename=self.cfg["file"])

            use_gpu = self.use_gpu_var.get()
            exe = SERVER_EXE_CUDA if use_gpu else SERVER_EXE_CPU
            ngl = str(self.cfg["gpu_layers"]) if use_gpu else "0"
            args = [exe, "-m", model_path, "--host", "127.0.0.1", "--port", str(self.port),
                    "-c", str(N_CTX), "-t", str(N_THREADS), "-ngl", ngl]

            if self.cfg["mmproj"]:
                self.events.put(("status", "Downloading vision encoder (first run only)…"))
                mmproj_path = hf_hub_download(repo_id=self.cfg["repo"], filename=self.cfg["mmproj"])
                args += ["--mmproj", mmproj_path]

            self.events.put(("status", f"Starting engine for Gemma 3 {self.current_size}…"))
            self._log_path = os.path.join(HERE, "server.log")
            log = open(self._log_path, "w", encoding="utf-8")
            self.proc = subprocess.Popen(args, stdout=log, stderr=subprocess.STDOUT,
                                         creationflags=CREATE_NO_WINDOW)

            deadline = time.time() + 240
            while time.time() < deadline:
                if self.proc.poll() is not None:
                    raise RuntimeError("engine exited early — see server.log")
                if self._health_ok():
                    if use_gpu and self._gpu_init_failed():
                        self.events.put(("gpu_fallback", None))
                    self.events.put(("ready", None))
                    return
                time.sleep(0.5)
            raise TimeoutError("engine did not become ready in time")
        except Exception as e:
            self.events.put(("error", f"{type(e).__name__}: {e}"))

    def _gpu_init_failed(self):
        try:
            with open(self._log_path, "r", encoding="utf-8", errors="ignore") as f:
                return "failed to initialize CUDA" in f.read()
        except Exception:
            return False

    def _stop_server(self):
        if self.proc and self.proc.poll() is None:
            self.proc.terminate()
            try:
                self.proc.wait(timeout=5)
            except Exception:
                self.proc.kill()
        self.proc = None

    def _health_ok(self):
        try:
            with urllib.request.urlopen(f"{self.base_url}/health", timeout=2) as r:
                return r.status == 200
        except Exception:
            return False

    # --- Sending / generation ----------------------------------------------
    def _on_send(self):
        text = self.entry.get().strip()
        if not text and not self.pending_attach:
            return
        self.entry.delete(0, "end")
        self._set_input_enabled(False)
        self.status.configure(text="Thinking…")

        first = self.store.message_count(self.current_chat_id) == 0
        att = self.pending_attach
        if att and att["kind"] == "image":
            content = self._image_content(text, att["path"])
            disp = text or "Describe this image."
            self._append("img", "You", f"   [🖼 {att['name']}] ")
            self._append("body", "", disp + "\n\n")
            self.store.add_message(self.current_chat_id, "user", text, att["path"])
            self._clear_attach()
        elif att and att["kind"] == "doc":
            content = self._doc_content(att["name"], att["text"], text)
            disp = text or "(review the attached document)"
            self._append("img", "You", f"   [📎 {att['name']}] ")
            self._append("body", "", disp + "\n\n")
            self.store.add_message(self.current_chat_id, "user", text, att["path"])
            self._clear_attach()
        else:
            content = text
            self._append("user", "You", "\n")
            self._append("body", "", text + "\n\n")
            self.store.add_message(self.current_chat_id, "user", text)

        self.messages.append({"role": "user", "content": content})
        if first:
            title = (text or (att["name"] if att else "")).strip()[:40] or "New conversation"
            self.store.set_title(self.current_chat_id, title)
            self.chat_title.configure(text=title)
            self._refresh_chat_list()

        self._append("ai", "Gemma", "\n")
        self._ai_body_index = self.txt.index("end-1c")   # remember where the reply starts
        threading.Thread(target=self._generate, daemon=True).start()

    def _sanitized_messages(self):
        """Gemma requires strictly alternating user/assistant turns. Merge any
        consecutive same-role messages (e.g. orphaned turns) so the request is valid."""
        out = []
        for m in self.messages:
            if out and out[-1]["role"] == m["role"] and m["role"] != "system":
                out[-1] = self._merge(out[-1], m)
            else:
                out.append(dict(m))
        return out

    @staticmethod
    def _merge(a, b):
        def parts(c):
            return list(c) if isinstance(c, list) else [{"type": "text", "text": c or ""}]
        merged = parts(a["content"]) + parts(b["content"])
        if all(p.get("type") == "text" for p in merged):
            return {"role": a["role"], "content": "\n".join(p["text"] for p in merged)}
        return {"role": a["role"], "content": merged}

    def _generate(self):
        try:
            body = json.dumps({
                "messages": self._sanitized_messages(), "stream": True,
                "temperature": 0.7, "max_tokens": 1024,
                "stream_options": {"include_usage": True},
            }).encode("utf-8")
            req = urllib.request.Request(f"{self.base_url}/v1/chat/completions",
                                         data=body, headers={"Content-Type": "application/json"})
            reply = ""
            usage = {}
            t_start = time.perf_counter()
            t_first = None
            with urllib.request.urlopen(req) as resp:
                for raw in resp:
                    line = raw.decode("utf-8").strip()
                    if not line.startswith("data: "):
                        continue
                    data = line[6:]
                    if data == "[DONE]":
                        break
                    obj = json.loads(data)
                    if obj.get("usage"):
                        usage = obj["usage"]
                    choices = obj.get("choices") or []
                    if choices:
                        delta = choices[0]["delta"].get("content")
                        if delta:
                            if t_first is None:
                                t_first = time.perf_counter()
                            reply += delta
                            self.events.put(("token", delta))
            t_end = time.perf_counter()
            self.messages.append({"role": "assistant", "content": reply})
            metrics = self._format_metrics(usage, t_start, t_first, t_end)
            self.events.put(("answer_done", (reply, metrics)))
        except Exception as e:
            self.events.put(("error", f"{type(e).__name__}: {e}"))

    @staticmethod
    def _format_metrics(usage, t_start, t_first, t_end):
        total = t_end - t_start
        in_tok = usage.get("prompt_tokens", 0)
        out_tok = usage.get("completion_tokens", 0)
        ttft = (t_first - t_start) if t_first else total
        gen_time = (t_end - t_first) if t_first else 0.0
        tps = (out_tok / gen_time) if gen_time > 0 else 0.0
        return (f"   ⏱ {total:.1f}s   ·   in {in_tok} tok   ·   out {out_tok} tok"
                f"   ·   {tps:.1f} tok/s   ·   first token {ttft:.1f}s\n\n")

    # --- Live resource monitor ---------------------------------------------
    def _monitor(self):
        try:
            pids = {os.getpid()}
            if self.proc and self.proc.poll() is None:
                pids.add(self.proc.pid)
                try:
                    for c in psutil.Process(self.proc.pid).children(recursive=True):
                        pids.add(c.pid)
                except psutil.Error:
                    pass
            for dead in [pid for pid in self._ps if pid not in pids]:
                self._ps.pop(dead, None)
            cpu = 0.0
            rss = 0
            for pid in pids:
                try:
                    p = self._ps.get(pid)
                    if p is None:
                        p = psutil.Process(pid)
                        p.cpu_percent(None)
                        self._ps[pid] = p
                    cpu += p.cpu_percent(None)
                    rss += p.memory_info().rss
                except psutil.Error:
                    self._ps.pop(pid, None)
            ncpu = psutil.cpu_count() or 1
            self.res_label.configure(text=f"CPU {cpu / ncpu:4.0f}%   RAM {rss / 1073741824:5.2f} GB")
        except Exception:
            pass
        self.root.after(1500, self._monitor)

    # --- Event pump --------------------------------------------------------
    def _drain_events(self):
        try:
            while True:
                kind, payload = self.events.get_nowait()
                if kind == "status":
                    self.status.configure(text=payload, text_color=YELLOW)
                elif kind == "ready":
                    self._ready_status()
                    self._set_input_enabled(True)
                elif kind == "token":
                    self._append("body", "", payload)
                elif kind == "answer_done":
                    reply, metrics = payload
                    if self.current_chat_id is not None:
                        self.store.add_message(self.current_chat_id, "assistant", reply)
                    self._finalize_reply(reply, metrics)
                    self._ready_status()
                    self._set_input_enabled(True)
                elif kind == "gpu_fallback":
                    self.use_gpu_var.set(False)
                    self._append("warn", "", "\n[GPU] CUDA could not initialize (driver too old) — "
                                             "running on CPU. Update the NVIDIA driver (≥551) and reboot.\n\n")
                elif kind == "error":
                    self._append("warn", "", f"\n[error] {payload}\n\n")
                    self.status.configure(text="Error — see chat / server.log", text_color=RED)
                    self._set_input_enabled(True)
        except queue.Empty:
            pass
        self.root.after(50, self._drain_events)

    def _ready_status(self):
        vis = "vision" if self.supports_vision else "text-only"
        backend = f"GPU (-ngl {self.cfg['gpu_layers']})" if self.use_gpu_var.get() else f"CPU · {N_THREADS} threads"
        self.status.configure(text=f"  Ready · Gemma 3 {self.current_size} ({vis}) · {backend} · :{self.port}",
                              text_color=GREEN)

    def _on_close(self):
        self._stop_server()
        self.root.destroy()


if __name__ == "__main__":
    root = ctk.CTk()
    ChatApp(root)
    root.mainloop()
